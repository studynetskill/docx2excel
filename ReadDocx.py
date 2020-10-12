from docx import Document  # 导入库


# 读取 path下的docx，并返回数组
class ReadDocx():
    '''
    读取 path下的docx，并返回数组
    '''

    def __init__(self, path):
        '''
        :param path: docx文件路径
        '''
        self.document = Document(path)  # 读入文件
        self.tables = self.document.tables  # 获取文件中的表格集

    def read_table(self):
        '''
        读取文件的table数据
        '''
        table_data = []
        cell_set = []
        for table in self.tables[:]:
            for i, row in enumerate(table.rows[:]):   # 读每行
                # row_content = []
                for cell in row.cells[:]:  # 读一行中的所有单元格
                    if cell not in cell_set:  # 同一行合并单元格的地址不变，合并的单元格虽然重复但公用内存地址
                        cell_set.append(cell)
                        c = cell.text
                        table_data.append(c.strip())
                # print(row_content)  # 以列表形式导出每一行数据
                # table_data.append(row_content)
        # print(table_data)
        return table_data


if __name__ == "__main__":
    path = '123.docx'  # 文件路径
    table = ReadDocx(path)
    table_data = table.read_table()
    print(table_data)
